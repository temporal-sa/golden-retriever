from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path("/Users/samingbar/Documents/Retrievals")
OUT = ROOT / "artifacts" / "repository-vs-original-design.docx"
BASELINE_IMAGE = Path("/Users/samingbar/Downloads/Whiteboard workflow DAG.png")
SKILL_SCRIPTS = Path(
    "/Users/samingbar/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.715.12143/skills/documents/scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402


NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF4D6"
PALE_RED = "FCE8E6"
WHITE = "FFFFFF"
BLACK = "1F1F1F"
BORDER = "C9D1DA"
GREEN = "236B4E"
GOLD = "7A5A00"
RED = "9B1C1C"


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, *, color: str = BORDER, size: int = 6) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_paragraph_shading(paragraph, fill: str, *, left_border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if left_border:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        left = p_bdr.find(qn("w:left"))
        if left is None:
            left = OxmlElement("w:left")
            p_bdr.append(left)
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), left_border)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_section(section, *, landscape: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Repository vs. Original Design")
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("18 July 2026  |  Page ")
    set_run_font(run, size=9, color=MUTED)
    add_page_number_field(p)


def set_style_font(style, name: str, size: float, color: str, *, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, "Calibri", 11, BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    h1 = styles["Heading 1"]
    set_style_font(h1, "Calibri", 16, BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True

    h2 = styles["Heading 2"]
    set_style_font(h2, "Calibri", 13, BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    h3 = styles["Heading 3"]
    set_style_font(h3, "Calibri", 12, DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.keep_together = True

    if "Evidence" not in styles:
        evidence = styles.add_style("Evidence", WD_STYLE_TYPE.PARAGRAPH)
    else:
        evidence = styles["Evidence"]
    set_style_font(evidence, "Consolas", 8.5, MUTED)
    evidence.paragraph_format.space_before = Pt(0)
    evidence.paragraph_format.space_after = Pt(3)
    evidence.paragraph_format.line_spacing = 1.0

    if "Table Body" not in styles:
        table_body = styles.add_style("Table Body", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_body = styles["Table Body"]
    set_style_font(table_body, "Calibri", 9.2, BLACK)
    table_body.paragraph_format.space_before = Pt(0)
    table_body.paragraph_format.space_after = Pt(3)
    table_body.paragraph_format.line_spacing = 1.05

    if "Table Header" not in styles:
        table_header = styles.add_style("Table Header", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_header = styles["Table Header"]
    set_style_font(table_header, "Calibri", 9.2, NAVY, bold=True)
    table_header.paragraph_format.space_before = Pt(0)
    table_header.paragraph_format.space_after = Pt(0)
    table_header.paragraph_format.line_spacing = 1.0


def add_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(x.get(qn("w:abstractNumId")))
        for x in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_base = max(abstract_ids, default=0) + 1
    num_base = max(num_ids, default=0) + 1

    def create(kind: str, abstract_id: int, num_id: int) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        abstract.append(lvl)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    create("bullet", abstract_base, num_base)
    create("decimal", abstract_base + 1, num_base + 1)
    return num_base, num_base + 1


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


def add_bullet(doc: Document, num_id: int, text: str, *, lead: str | None = None) -> None:
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    if lead:
        r = p.add_run(lead)
        set_run_font(r, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run_font(r)


def add_labeled_paragraph(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"{label}: ")
    set_run_font(r, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run_font(r)


def add_callout(doc: Document, label: str, text: str, *, fill: str = PALE_BLUE) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.10)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.10
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, fill, left_border=BLUE)
    r = p.add_run(f"{label}  ")
    set_run_font(r, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run_font(r, color=BLACK)


def add_status_line(doc: Document, status: str, text: str, color: str, fill: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, fill)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    r = p.add_run(f"{status.upper()}  ")
    set_run_font(r, size=9.5, color=color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.2, color=BLACK)


def format_table(table, *, header_fill: str = LIGHT_GRAY) -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
            for p in cell.paragraphs:
                p.style = "Table Header" if row_index == 0 else "Table Body"
                p.paragraph_format.keep_together = True
                for run in p.runs:
                    set_run_font(
                        run,
                        size=9.2,
                        color=NAVY if row_index == 0 else BLACK,
                        bold=row_index == 0,
                    )
    set_repeat_table_header(table.rows[0])


def add_table_spacer(doc: Document, *, before: float = 4, after: float = 7) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1


def add_comparison_table(doc: Document) -> None:
    rows = [
        (
            "Control plane",
            "RootSyncWorkflow is the visible top-level owner; deactivation is initiated from the root.",
            "StoreControllerWorkflow serializes commands and starts RootSyncWorkflow and DeactivateStoreWorkflow as detached, stable-ID operations.",
        ),
        (
            "Sync fan-out",
            "Page barriers, round windows, joined resources, sliding page children, and joined document work.",
            "Core shape retained; every major fan-out is bounded, cancellation-aware, checkpointed, and given deterministic child IDs.",
        ),
        (
            "Quota",
            "A same-namespace QuotaWaitWorkflow child is awaited; it runs a heartbeat Activity until signaled.",
            "Callers submit to one shared UserQuotaWorkflow per quota scope, wait on Signals in their own history, and explicitly complete or cancel permits.",
        ),
        (
            "Remediation",
            "FailedUserRemediationWorkflow is detached after child-start acknowledgement and its handle is discarded.",
            "Still detached, but the root caps active remediation starts and registers each with the controller, which owns status until terminal.",
        ),
        (
            "Deactivation",
            "The root awaits only a newly started child; duplicate-start is swallowed without joining the existing execution.",
            "The controller returns the stable operation identity; deactivation runs fence -> cancel -> quota invalidation -> bounded drain -> cleanup -> inactive.",
        ),
        (
            "Resource routing",
            "Resource-family activities, a comments follow-up, and direct document-child conversion are distinct branches.",
            "A unified provider page-manifest Activity feeds FilesPageWorkflow; CommentsResyncWorkflow remains only as a direct wrapper boundary.",
        ),
        (
            "Legacy workflows",
            "QuotaWaitWorkflow, AccessioningWorkflow, and a separate rate-limit workflow are active design elements.",
            "QuotaWaitWorkflow and AccessioningWorkflow are optional drain-only names; the separate rate-limit workflow is absent from the primary registry.",
        ),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Area"
    table.rows[0].cells[1].text = "Original design"
    table.rows[0].cells[2].text = "Current repository"
    for area, original, current in rows:
        cells = table.add_row().cells
        cells[0].text = area
        cells[1].text = original
        cells[2].text = current
    apply_table_geometry(table, [1500, 3420, 4440], table_width_dxa=9360, indent_dxa=120)
    format_table(table)


def add_inventory_table(doc: Document) -> None:
    rows = [
        (
            "Added",
            "StoreControllerWorkflow; UserQuotaWorkflow",
            "New workflow histories, Signals, stable external ownership, and independent Continue-As-New policies.",
        ),
        (
            "Retained in primary path",
            "RootSync, failed-user remediation, activation, user/resource/page/file/document spine, cleanup workflows",
            "Original boundaries remain recognizable, but carry new inputs, IDs, cancellation policies, and generation state.",
        ),
        (
            "Consolidated",
            "Resource-family page branches",
            "provider_fetch_resource_page returns a compact manifest; FilesPageWorkflow handles upserts and deletions uniformly.",
        ),
        (
            "Direct-only",
            "CommentsResyncWorkflow",
            "Registered for explicit callers, but the controller-driven sync tree does not start it.",
        ),
        (
            "Legacy-only / absent",
            "QuotaWaitWorkflow; AccessioningWorkflow; separate rate-limit workflow",
            "First two are optional drain placeholders, not replay-compatible replacements; the third is not registered.",
        ),
        (
            "New Activity boundary",
            "StagingStore-backed ingestion",
            "Only DocumentRef metadata enters Workflow History; the ingestion Activity loads the body and commits a generation-fenced mutation.",
        ),
        (
            "New queue boundary",
            "Provider Task Queue",
            "Provider API Activities are isolated from workflows and persistence Activities and can receive queue rate limits and priority metadata.",
        ),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Delta"
    table.rows[0].cells[1].text = "Types / boundary"
    table.rows[0].cells[2].text = "Runtime meaning"
    for delta, boundary, meaning in rows:
        cells = table.add_row().cells
        cells[0].text = delta
        cells[1].text = boundary
        cells[2].text = meaning
    apply_table_geometry(table, [1600, 3100, 4660], table_width_dxa=9360, indent_dxa=120)
    format_table(table)


def add_evidence_table(doc: Document) -> None:
    rows = [
        ("Controller ownership and duplicate commands", "src/retrieval/temporal/workflows/store_controller.py:284-376, 411-536"),
        ("Root page/round barriers and remediation registration", "src/retrieval/temporal/workflows/root_sync.py:242-272, 279-413, 501-680, 710-802"),
        ("Resource fan-out and activation waves", "src/retrieval/temporal/workflows/user_sync.py:30-136; activate_user.py:60-148"),
        ("Page sliding window, checkpoint, and rollover", "src/retrieval/temporal/workflows/resource_pages.py:57-130, 162-282"),
        ("Document fan-out and staged mutation", "src/retrieval/temporal/workflows/files_page.py:34-123; document_ingestion.py:24-34; activities/ingestion.py:66-125"),
        ("Shared quota protocol", "src/retrieval/temporal/common/quota_waiter.py:68-326; workflows/user_quota.py:557-813; activities/quota_client.py:25-75"),
        ("Legacy drain names", "src/retrieval/temporal/workflows/legacy.py:1-26; worker.py:62-80, 119-122"),
        ("Fence/cancel/drain/cleanup deactivation", "src/retrieval/temporal/workflows/deactivate_store.py:183-319, 367-484; cleanup.py:150-230"),
        ("Worker and Task Queue split", "src/retrieval/temporal/worker.py:98-149"),
        ("Repository architecture statements", "README.md; IMPLEMENTATION_MAP.md; docs/workflow-topology.md; docs/adr/0001-workflow-boundaries.md"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Claim area"
    table.rows[0].cells[1].text = "Current-working-tree evidence"
    for claim, evidence in rows:
        cells = table.add_row().cells
        cells[0].text = claim
        cells[1].text = evidence
    apply_table_geometry(table, [2850, 6510], table_width_dxa=9360, indent_dxa=120)
    format_table(table)
    for row in table.rows[1:]:
        for run in row.cells[1].paragraphs[0].runs:
            set_run_font(run, name="Consolas", size=8.2, color=MUTED)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("ARCHITECTURE COMPARISON")
    set_run_font(r, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Temporal Retrieval Workflows")
    set_run_font(r, size=25, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Current repository vs. the original whiteboard design")
    set_run_font(r, size=13.5, color=MUTED)

    metadata = [
        ("Baseline", "Provided Whiteboard workflow DAG plus the eight execution characteristics supplied with the request"),
        ("Repository snapshot", "Current on-disk working tree, 18 July 2026"),
        ("Validation", "65 targeted Temporal semantic tests passed"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_run_font(r, size=9.5, color=NAVY, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=9.5, color=MUTED)


def build_document() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_section(doc.sections[0])
    configure_header_footer(doc.sections[0])
    configure_styles(doc)
    bullet_num_id, _decimal_num_id = add_numbering(doc)

    props = doc.core_properties
    props.title = "Temporal Retrieval Workflows: Current Repository vs. Original Design"
    props.subject = "Differences in Temporal workflow topology and execution semantics"
    props.author = "Codex"
    props.keywords = "Temporal, workflows, architecture, comparison, retrieval"
    props.comments = "Generated from the current working tree and the supplied original design baseline."

    add_title_block(doc)
    add_callout(
        doc,
        "BOTTOM LINE",
        "The repository is an evolved architecture, not a literal transcription of the whiteboard. "
        "It preserves the central joined sync spine and the original page/round barriers, while "
        "changing lifecycle ownership, quota coordination, deactivation safety, resource-page "
        "routing, payload handling, and operational controls.",
    )

    doc.add_heading("Executive readout", level=1)
    add_bullet(doc, bullet_num_id, " Root user-page barriers, round-mode carry-forward, joined resource fan-out, the sliding FilesPage window, and two-wave activation remain recognizable.", lead="Preserved.")
    add_bullet(doc, bullet_num_id, " QuotaWaitWorkflow is no longer the waiting mechanism; a shared UserQuotaWorkflow coordinates permits for each provider/credential/quota-class scope.", lead="Replaced.")
    add_bullet(doc, bullet_num_id, " Failed-user remediation is still detached, but it is no longer intentionally unowned after start; it is registered with and tracked by the store controller.", lead="Ownership changed.")
    add_bullet(doc, bullet_num_id, " Deactivation is no longer a root-child special case. It is a controller-started stable operation with a generation fence and an explicit fence -> cancel -> drain -> cleanup sequence.", lead="Safety boundary changed.")
    add_bullet(doc, bullet_num_id, " Concurrency ceilings, deterministic IDs, cursor checkpoints, bounded result samples, explicit cancellation behavior, and additional Continue-As-New boundaries are now first-class policies.", lead="Hardened.")

    doc.add_page_break()
    doc.add_heading("At-a-glance comparison", level=1)
    p = doc.add_paragraph(
        "“Retained” below means the original Temporal completion barrier still exists. “Hardened” "
        "means the shape remains but policy and state were added. “Replaced” means the Event History "
        "and integration contract materially differ."
    )
    p.paragraph_format.space_after = Pt(7)
    add_comparison_table(doc)
    add_table_spacer(doc)

    doc.add_page_break()
    doc.add_heading("1. Control-plane ownership is new", level=1)
    add_status_line(
        doc,
        "Replaced",
        "The current entry point is StoreControllerWorkflow, not RootSyncWorkflow.",
        RED,
        PALE_RED,
    )
    add_labeled_paragraph(
        doc,
        "Original design",
        "The visible graph begins at RootSyncWorkflow. The root owns normal synchronization, starts detached remediation, and also contains the special deactivation-start behavior.",
    )
    add_labeled_paragraph(
        doc,
        "Current repository",
        "Applications submit idempotent Updates through RetrievalClient to one long-lived StoreControllerWorkflow per store. The controller serializes sync, cancel, and deactivation commands; enforces generation and lifecycle rules; starts RootSyncWorkflow and DeactivateStoreWorkflow with ParentClosePolicy.ABANDON; and tracks terminal status through Signals.",
    )
    add_labeled_paragraph(
        doc,
        "Temporal consequence",
        "High-volume sync history is separated from low-volume lifecycle authority. Duplicate commands return a stable operation identity instead of relying on child-start exceptions, and the store remains SYNCING until both root work and detached remediation are terminal.",
    )
    add_callout(
        doc,
        "INTEGRATION CHANGE",
        "Callers should use controller Updates and the controller status query. Directly starting the root or inferring deactivation completion from duplicate-start behavior bypasses the repository’s authority model.",
        fill=PALE_GOLD,
    )

    doc.add_heading("2. Original execution characteristics, one by one", level=1)

    doc.add_heading("2.1 Root ordinary mode: retained, then bounded", level=2)
    add_status_line(
        doc,
        "Retained + hardened",
        "Each user page is still a completion barrier using gather(..., return_exceptions=True).",
        GREEN,
        PALE_BLUE,
    )
    add_labeled_paragraph(
        doc,
        "What stayed the same",
        "The root lists active users by page, filters invalid users, creates one coroutine per valid user, and does not advance the user cursor until the batch has joined. Per-user failures are classified after the gather rather than failing the whole page immediately.",
    )
    add_labeled_paragraph(
        doc,
        "What changed",
        "A semaphore caps concurrent UserSyncWorkflow starts at max_active_users. Children receive deterministic IDs and WAIT_CANCELLATION_COMPLETED semantics. Cancellation explicitly cancels local tasks and active child handles. Ordinary mode also carries cumulative counters and bounded samples across page-boundary Continue-As-New.",
    )

    doc.add_heading("2.2 Connector round mode: retained with fuller carry state", level=2)
    add_status_line(
        doc,
        "Retained + hardened",
        "The bounded active-user window, per-user page slice, whole-round join, and carry-forward model remain.",
        GREEN,
        PALE_BLUE,
    )
    add_labeled_paragraph(
        doc,
        "Current refinements",
        "The carried RoundState includes active users, buffered users, next user cursor, round number, exhaustion state, per-resource cursors, completed resource types, and bounded failure samples. Continue-As-New occurs only after the round is drained; the root also drains the remediation-handle set before rolling over.",
    )
    add_labeled_paragraph(
        doc,
        "Temporal consequence",
        "The fairness model remains round-based, but rollover is now an explicit safe barrier with enough state to resume unfinished users without replaying completed resource types.",
    )

    doc.add_heading("2.3 User fan-out and activation waves: preserved with a fence", level=2)
    add_status_line(
        doc,
        "Retained + hardened",
        "UserSyncWorkflow still joins one ResourceSyncWorkflow per requested resource, and activation still runs recent data before full backfill.",
        GREEN,
        PALE_BLUE,
    )
    add_labeled_paragraph(
        doc,
        "What changed",
        "Resource fan-out is bounded by resource_concurrency and skips resource types already completed in carried state. After the capped recent wave, ActivateUserWorkflow performs a lifecycle-generation validation Activity before launching backfill, then performs a generation-fenced activation mutation after backfill succeeds.",
    )
    add_labeled_paragraph(
        doc,
        "Temporal consequence",
        "The two waves remain sequential, but deactivation or another generation change can stop the second wave and the final user-state mutation even when the recent child completed successfully.",
    )

    doc.add_heading("2.4 Resource page window: retained with failure checkpointing", level=2)
    add_status_line(
        doc,
        "Retained + hardened",
        "ResourcePagesWorkflow still fills a sliding FilesPageWorkflow window and waits for FIRST_COMPLETED when full.",
        GREEN,
        PALE_BLUE,
    )
    add_labeled_paragraph(
        doc,
        "What stayed the same",
        "Started page children drain before return, cancellation, or Continue-As-New. The next page is admitted only while the configured window has room.",
    )
    add_labeled_paragraph(
        doc,
        "What changed",
        "The workflow obtains a shared quota permit before each provider-fetch Activity, stops admitting new pages after a failure, drains the already-started set, and checkpoints the earliest failed page’s input cursor. A deterministic attempt suffix prevents reuse of a terminal permit request when quota exhaustion is retried without advancing the cursor.",
    )

    doc.add_heading("2.5 FilesPage and document ingestion: expanded and generation-fenced", level=2)
    add_status_line(
        doc,
        "Changed at the edge",
        "Document children remain joined, but the current path handles both upserts and deletes through DocumentIngestionWorkflow.",
        GOLD,
        PALE_GOLD,
    )
    add_labeled_paragraph(
        doc,
        "Current repository",
        "FilesPageWorkflow creates one task for each document upsert and each deleted document key, bounds child execution with the lower configured ingestion ceiling, and uses a deterministic ID derived from store, generation, document key, and source version. The child runs a staged-document Activity; document bodies stay outside Workflow History.",
    )
    add_labeled_paragraph(
        doc,
        "Temporal consequence",
        "Every mutation has a durable child boundary and atomic lifecycle-generation check. Already-started work after an earlier page failure may run again, so stable IDs and idempotent repository writes are required rather than merely desirable.",
    )

    doc.add_page_break()
    doc.add_heading("2.6 Quota waiting: fully replaced", level=2)
    add_status_line(
        doc,
        "Replaced",
        "QuotaWaitWorkflow is not started by the current execution path.",
        RED,
        PALE_RED,
    )
    add_labeled_paragraph(
        doc,
        "Original design",
        "A caller executes and awaits a same-namespace QuotaWaitWorkflow child. That child repeatedly runs a heartbeat Activity until a matching signal arrives; separate accessioning/rate-limit workflows participate outside the retrieval namespace.",
    )
    add_labeled_paragraph(
        doc,
        "Current repository",
        "RootSyncWorkflow and ResourcePagesWorkflow call a reusable quota-waiter mixin. A short Activity performs Signal-with-Start against one UserQuotaWorkflow per provider, opaque credential key, and quota class. The caller installs an expected-request inbox before that Activity, then waits durably on quota_granted or quota_denied Signals in its own workflow. It later sends permit_completed, cancel_permit, or observe_quota as appropriate.",
    )
    add_labeled_paragraph(
        doc,
        "Coordinator behavior",
        "The shared workflow owns the pending queue, in-flight reservations, reset timer, observations, deduplication, cancellation by lifecycle generation, explicit denials, and its own Continue-As-New thresholds. Waiting consumes no Activity worker slot.",
    )
    add_labeled_paragraph(
        doc,
        "Compatibility warning",
        "QuotaWaitWorkflow and AccessioningWorkflow exist only as optional drain-name placeholders. They cannot replay arbitrary histories produced by their original implementations. Open legacy histories must remain pinned to a compatible worker build.",
    )

    doc.add_heading("2.7 Failed-user remediation: still detached, now durably owned", level=2)
    add_status_line(
        doc,
        "Ownership changed",
        "The ABANDON relationship remains, but the handle is not simply discarded after start.",
        GOLD,
        PALE_GOLD,
    )
    add_labeled_paragraph(
        doc,
        "Current repository",
        "The root awaits ChildWorkflowExecutionStarted, then shields registration of the remediation with StoreControllerWorkflow. If registration fails, it cancels the child rather than leaving it unowned. The root holds a bounded active-remediation set, starts at most four concurrently, and drains one with FIRST_COMPLETED before exceeding the cap. Remediation itself batches activation children and may Continue-As-New between batches.",
    )
    add_labeled_paragraph(
        doc,
        "Temporal consequence",
        "RootSyncWorkflow may still finish before remediation, but the controller keeps the store in SYNCING, rejects another sync, cancels late remediation during deactivation, and receives an idempotent terminal signal.",
    )

    doc.add_heading("2.8 Deactivation: ownership and completion semantics replaced", level=2)
    add_status_line(
        doc,
        "Replaced",
        "The root no longer owns or starts deactivation.",
        RED,
        PALE_RED,
    )
    add_labeled_paragraph(
        doc,
        "Original design",
        "The root waits for a newly started DeactivateStoreWorkflow. If a child with the same ID is already running, duplicate-start is swallowed and the root returns without joining the existing execution.",
    )
    add_labeled_paragraph(
        doc,
        "Current repository",
        "StoreControllerWorkflow creates a generation-derived stable operation and returns Accepted. Repeated command IDs and already-active deactivation requests return that same operation identity. The detached deactivation reports fenced and terminal states to the controller rather than being joined by the caller.",
    )
    add_labeled_paragraph(
        doc,
        "New safety sequence",
        "DeactivateStoreWorkflow first commits the lifecycle-generation fence. Only then does it request cancellation of tracked sync/remediation, invalidate old-generation quota requests, wait for a bounded drain, clean users, remove objects, mark the store inactive, and signal terminal status. Once fenced, its protected cleanup continues toward a safe terminal state even if outer cancellation arrives.",
    )
    add_labeled_paragraph(
        doc,
        "Cleanup topology",
        "The former side branches are ordered. DeactivateStoreWorkflow executes one CleanupUsersWorkflow, which chooses all-user cleanup for an empty set or bounded per-user batches for explicit keys; only after user cleanup succeeds does RemoveObjectsWorkflow run.",
    )

    doc.add_heading("3. Topology changes beyond the eight characteristics", level=1)
    add_inventory_table(doc)
    add_table_spacer(doc)

    doc.add_page_break()
    doc.add_heading("4. Why these differences matter to Temporal", level=1)
    add_bullet(doc, bullet_num_id, " The controller, detached status Signals, shared quota workflow, and reordered deactivation produce materially different command sequences and histories. These are architecture changes, not implementation-detail substitutions.", lead="Replay compatibility.")
    add_bullet(doc, bullet_num_id, " Stable IDs and explicit ParentClosePolicy / cancellation types make ownership intentional. ABANDON now means “detached but externally tracked,” while joined children use WAIT_CANCELLATION_COMPLETED.", lead="Ownership.")
    add_bullet(doc, bullet_num_id, " Root ordinary pages, root rounds, page windows, remediation batches, the quota coordinator, and the idle controller have safe Continue-As-New boundaries. Each boundary carries only bounded state and occurs after the relevant child barrier.", lead="Event History.")
    add_bullet(doc, bullet_num_id, " The committed lifecycle generation, checked in the same transaction as every write, is the correctness barrier. Cancellation limits wasted work but is not relied on to stop late Activity completion.", lead="Late delivery safety.")
    add_bullet(doc, bullet_num_id, " Page-window success after an earlier failure can execute again from the earliest failed cursor. Staged references, deterministic IDs, and idempotent repository adapters are required for safe re-execution.", lead="Idempotency.")
    add_bullet(doc, bullet_num_id, " Provider API Activities run on a separate Task Queue. Shared quota admission happens before those Activities, and optional Temporal priority/fairness metadata applies after admission.", lead="Scheduling.")

    doc.add_heading("5. Practical migration and review implications", level=1)
    add_bullet(doc, bullet_num_id, " Use RetrievalClient and StoreControllerWorkflow Updates as the public command surface; do not preserve the original direct-root/deactivation call assumptions.")
    add_bullet(doc, bullet_num_id, " Keep compatible workers for open original QuotaWaitWorkflow or AccessioningWorkflow histories. The drain-only placeholders are names, not behavioral replacements.")
    add_bullet(doc, bullet_num_id, " Replay representative long-running, signaled, canceled, failed, retried, and Continue-As-New histories for every workflow type routed to a new build.")
    add_bullet(doc, bullet_num_id, " Update provider adapters to return compact page manifests and stage bodies outside Workflow History; map quota exhaustion to structured reset observations.")
    add_bullet(doc, bullet_num_id, " Implement production persistence with atomic generation/status compare-and-write and idempotent document/cleanup mutations. The bundled in-memory repository, staging store, and empty provider are development adapters only.")
    add_callout(
        doc,
        "REPOSITORY SCOPE",
        "The working tree is a reference implementation for local development and architecture evaluation, not a deployable production service by itself. Production adapters, telemetry export, representative history replay, target-namespace validation, and capacity evidence remain external release gates.",
        fill=PALE_GOLD,
    )

    doc.add_heading("6. Evidence and scope", level=1)
    add_labeled_paragraph(
        doc,
        "Snapshot basis",
        "This comparison describes the current files on disk on 18 July 2026. The repository had uncommitted modifications, so the document intentionally does not claim to describe only commit 35433be.",
    )
    add_labeled_paragraph(
        doc,
        "Original-design basis",
        "The supplied Whiteboard workflow DAG and the eight execution characteristics in the request. No separate original source repository was provided, so the baseline is interpreted from those materials rather than reconstructed from an older code revision.",
    )
    add_labeled_paragraph(
        doc,
        "Validation",
        "The targeted suites for sync invariants, deactivation order, controller contract, and quota logic completed successfully: 65 passed in 0.22 seconds.",
    )
    add_evidence_table(doc)
    add_table_spacer(doc)

    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(landscape, landscape=True)
    # Header/footer remain linked to the first section. Do not append duplicate content.
    appendix_heading = doc.add_heading("Appendix A. Original design baseline", level=1)
    appendix_heading.paragraph_format.space_before = Pt(0)
    appendix_heading.paragraph_format.space_after = Pt(4)
    p = doc.add_paragraph("Provided whiteboard used as the comparison baseline.")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for caption_run in p.runs:
        set_run_font(caption_run, size=9.5, color=MUTED, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(BASELINE_IMAGE), width=Inches(7.25))
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    inline = doc.inline_shapes[-1]._inline
    inline.docPr.set("descr", "Original Temporal retrieval workflow DAG provided by the user")
    inline.docPr.set("title", "Original design baseline")

    doc.save(OUT)


if __name__ == "__main__":
    build_document()
